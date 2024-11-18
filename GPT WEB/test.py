from docx import Document

doc = Document("C:/Users/qwe/Desktop/VSCODE/GPT WEB/docx_storage/GPT 고소장_문문문_계약금사기_025408.docx")

#table_receiver = doc.tables[1]
#print(doc.tables[1].rows[0].cells[1].paragraphs[0].text)
'''
for row in table_receiver.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            print(para.text)
'''

filename = "GPT 고소장_문문문_계약금사기_025408.docx"
print(filename.split("_")[2])